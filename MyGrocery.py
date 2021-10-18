#! /usr/bin/python
# -*- coding: utf-8 -*-

import sys, os
import re 
import uuid
import time
import datetime

import sqlite3
from PyQt5.QtSql import *

from PyQt5 import QtCore, QtGui, uic
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *

import docx

from docx import Document

#============================================================================================================DATA BASE
conn = sqlite3.connect('db.db')
# os.system("icacls db.db /grant *S-1-1-0:(D,WDAC)")	
query = conn.cursor()
conn.text_factory = str		
#===========================================================================TABLE LOGIN
try:
	# query.execute("DROP TABLE Login")
	query.execute("SELECT id FROM Login ORDER BY id DESC")
except:
	conn.execute("""CREATE TABLE Login (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					user_name VARCHAR(50) ,
					user_password VARCHAR(30))""")
					
	# print("DataBase  Login created succefully")
	query.execute("INSERT INTO Login (user_name, user_password) VALUES ('admin','admin')")
#===========================================================================TABLE User
try:
	# query.execute("DROP TABLE User")
	query.execute("SELECT user FROM User ")
except:
	# query.execute("DROP TABLE User")
	conn.execute("""CREATE TABLE User (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					user VARCHAR(50))""")
					
	# print("DataBase  User created succefully")
	query.execute("INSERT INTO User (user) VALUES ('')")
	
#===========================================================================TABLE SOCIETY
try:
	# query.execute("DROP TABLE Society")
	query.execute("SELECT id FROM Society ORDER BY id DESC")
except:
	conn.execute("""CREATE TABLE Society (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					society_name VARCHAR(50) ,
					society_number VARCHAR(30))""")
					
	# print("DataBase  Society created succefully")
	query.execute("INSERT INTO Society(society_name, society_number) VALUES ('Mon Entreprise','0549484715')")
#===========================================================================TABLE CATEGORIES
try:
	query.execute("SELECT id FROM Categories ORDER BY id DESC")
except:
	# query.execute("DROP TABLE Categories")
	conn.execute("""CREATE TABLE Categories (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					category_name VARCHAR(50))""")

#===========================================================================TABLE PRODUCTS
# query.execute("DROP TABLE Products")
try:
	query.execute("SELECT id FROM Products ORDER BY id DESC")

except:
	# query.execute("DROP TABLE Products")
	conn.execute("""CREATE TABLE Products (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					product_category VARCHAR(50),
					product_name VARCHAR(50) ,
					product_price INTEGER ,
					product_stock INTEGER ,
					product_stock_min INTEGER ,
					product_date DATE ,
					product_BC VARCHAR(20))""")

	
#===========================================================================TABLE PRODUCTS RECEIPT
# query.execute("DROP TABLE Receipt")
try:
	query.execute("SELECT id FROM Receipt ORDER BY id DESC")

except:
	# query.execute("DROP TABLE Receipt")
	conn.execute("""CREATE TABLE Receipt (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE,
					product_name VARCHAR(50) ,
					product_price INTEGER ,
					product_qnt INTEGER ,
					product_date DATE,
					product_ref INTEGER)""")
		
#===========================================================================REGISTER CATEGORY
# query.execute("DROP TABLE Register")
try:
	query.execute("SELECT id FROM Register ORDER BY id DESC")
except:
	conn.execute("""CREATE TABLE Register (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					register_date DATE,
					register_sum_init INTEGER,
					register_recette_total INTEGER,
					register_depense_total INTEGER,
					register_ajout_total INTEGER )""")	
#============================================================= REFERENCEMENT TICKETS
# query.execute("DROP TABLE Ref")
#===========================================================================DEPENSE 
# query.execute("DROP TABLE Register_dep")
try:
	query.execute("SELECT id FROM Register_dep ORDER BY id DESC")
except:
	conn.execute("""CREATE TABLE Register_dep (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					dep_date DATE,
					dep_type VARCHAR(15),
					dep_description VARCHAR(255),
					dep_value INTEGER)""")
					
	# print("DataBase  Register_dep created succefully")
					
#===========================================================================AJOUT 
# query.execute("DROP TABLE Register_add")
try:
	query.execute("SELECT id FROM Register_add ORDER BY id DESC")
except:
	conn.execute("""CREATE TABLE Register_add (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					add_date DATE,
					add_depo VARCHAR(50),
					add_description VARCHAR(255),
					add_value INTEGER)""")
					
	# print("DataBase  Register_add created succefully")
					
#===========================================================================REMISE 
# query.execute("DROP TABLE Discount")
try:
	query.execute("SELECT id FROM Discount ORDER BY id DESC")
except:
	conn.execute("""CREATE TABLE Discount (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					disc_date DATE,
					disc_ref INTEGER,
					disc_value INTEGER)""")
					
try:
	query.execute("SELECT id FROM Ref ORDER BY id DESC")
except:
	conn.execute("""CREATE TABLE Ref (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE, 
					ref INTEGER )""")
					
	query.execute("INSERT INTO Ref (ref) VALUES (0)")
	query.execute("UPDATE Ref SET ref = 0 WHERE id = 1")
	query.execute("UPDATE Ref SET ref = 0 WHERE id = 2")

#============================================================= MAC ADRESS
# query.execute("DROP TABLE MAC")
try:
	query.execute("SELECT id FROM MAC ORDER BY id DESC")
except:
	query.execute("""CREATE TABLE MAC (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE,
					Key VARCHAR(50))""")
	# print("MAC CREATED")


	query.execute("INSERT INTO MAC (Key) VALUES ('')")
#============================================================================================================END DATA BASE
conn.commit()

#======================================================CLASSES
#============================================================================================================MESSAGE FACTORY
class MessageFactory():

	def raiseNoPack(self):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Information)
		msg.setWindowTitle("Pack incompatible !")
		msg.setText("Votre pack ne contient pas cette fonctionnalité !\nContactez votre fournisseur pour plus d'infos.")
		msg.exec_()	

	def raiseNoConnect(self):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Information)
		msg.setWindowTitle("Connexiont impossible !")
		msg.setText("Vérifiez que les identifiants utilisés sont correct\nOu référer vous à votre fournisseur de produit!")
		msg.exec_()	

	def raiseNoMatch(self,data):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Information)
		msg.setWindowTitle("NO MATCH !")
		msg.setText(data+" est introuvable ou inéxistant !")
		msg.exec_()	
		
	def raiseMatch(self,data):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Information)
		msg.setWindowTitle("Identifiant existant !")
		msg.setText("l'utilisateur '"+data+"' est dêja pris !")
		msg.exec_()	
		
	def raiseNoInt(self):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Information)
		msg.setWindowTitle("Caractere non permis !")
		msg.setText(" Veuillez saisir un chiffre entier non nul!")
		msg.exec_()	
		
	def raiseAdder(self,data):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Information)
		msg.setWindowTitle("OPERATION REUSSIE !")
		msg.setText(data+" ajouté(e) avec succés !")
		msg.exec_()	
		
	def raiseModifier(self,data):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Information)
		msg.setWindowTitle("OPERATION REUSSIE !")
		msg.setText(data+" modifié(e) avec succés !")
		msg.exec_()	
		
	def raiseDeleter(self,data):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Information)
		msg.setWindowTitle("OPERATION REUSSIE !")
		msg.setText(data+" suprimé(e) avec succés !")
		msg.exec_()	
		
	def raiseCaseExcept(self,case):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setWindowTitle("CASE OBLIGATOIRE VIDE !")
		msg.setText("Veuillez saisir la case "+case)
		msg.exec_()
		
	def raisePrintExcept(self,doc):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setWindowTitle("ECHEC IMPRESSION !")
		msg.setText("Veuillez fermer le document word : '"+str(doc)+".docx'  !")
		msg.exec_()	
		
	def raiseCharExcept(self):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setWindowTitle("CARACTERE INDESIRABLE !")
		msg.setText("Veuillez ne pas utiliser de caracteres spéciaux ?;:',.$*.... !")
		msg.exec_()
		
	def raiseIndefinedExcept(self,erreur):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setWindowTitle("Warning !")
		msg.setText(str(erreur))
		msg.exec_()
		
	def raiseStockAlert(self,stock):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setWindowTitle("ALERT STOCK !")
		msg.setText("Vérifié votre stock de "+stock)
		msg.exec_()

class Selecter() :	
	def SelectOne(self,table, selection):
		lister = []
		query.execute("SELECT '{0}' FROM '{1}' ORDER BY id ASC".format(str(selection),str(table)))
		lister = query.fetchall()
		i=0
		for item in lister :
			lister[i] = str(item).strip("(',')")
			i+=1
		
		return lister
		
	def SelectOneCond(self,table, selection, condition, conditionInit):
		lister = []
		query.execute("SELECT '{0}' FROM '{1}' WHERE '{2}' = '{3}' ORDER BY id ASC".format(str(selection),str(table),str(condition),str(conditionInit)))
		lister = query.fetchall()
		i=0
		for item in lister :
			lister[i] = str(item).strip("(',')")
			i+=1
		
		return lister
		
class ReferenceSelector():
	def __init__(self):
		self.refSelector()
		self.societySelctor()
	
	def refSelector(self):
		query.execute("SELECT ref FROM Ref WHERE id=1")
		self.ref = str(query.fetchone()).strip("(',')")	
		
	def societySelctor(self):
		query.execute("SELECT society_name FROM Society")
		self.society = query.fetchone() 
		self.society = str(self.society).strip("(',')") 
		query.execute("SELECT society_number FROM Society")
		self.numero = query.fetchone() 
		self.numero = str(self.numero).strip("(',')") 	
#============================================================================================================DEV
qtValid= "DESIGN/DIALOGS/developper.ui"
Ui_devDialog, QtBaseClass = uic.loadUiType(qtValid)
class DevDialog(QDialog, Ui_devDialog):#CONFIRM : VALID DIALOG ON VLDButton CLICK()
	def __init__(self):
		QDialog.__init__(self)
		Ui_devDialog.__init__(self)
		self.setupUi(self)
		
		
#============================================================================================================SETTINGS
qtValid= "DESIGN/DIALOGS/settingCreatorDialog.ui"
Ui_SettingCreatorDialog, QtBaseClass = uic.loadUiType(qtValid)
class SettingCreatorDialog(QDialog, Ui_SettingCreatorDialog):#CONFIRM : VALID DIALOG ON VLDButton CLICK()
	def __init__(self):
		QDialog.__init__(self)
		Ui_SettingCreatorDialog.__init__(self)
		self.setupUi(self)
		
		self.referenceSelector = ReferenceSelector()
		self.messageFactory = MessageFactory()
		self.Selecter = Selecter()
		
		self.selection()
		self.settingSignal()
		
	def selection(self):
		self.society = self.referenceSelector.society 
		self.numero = self.referenceSelector.numero  
		
		self.societyName.setText(self.society)
		self.societyNumber.setText(self.numero)
		
		query.execute("SELECT user FROM User WHERE id = 1")
		self.user = query.fetchone()
		self.user= str(self.user).strip("(',')")
		self.who.setText(self.user)
		
		query.execute("SELECT user_name FROM Login ORDER BY id ASC")
		self.users=[]
		self.users = query.fetchall()
		query.execute("SELECT user_password FROM Login ORDER BY id ASC")
		self.passwords=[]
		self.passwords=query.fetchall()
		
		self.modSetComboBox.clear()
		self.modSetComboBox.addItem("...")
		self.delSetComboBox.clear()
		self.delSetComboBox.addItem("...")
		
		i=0
		for user in self.users :
			i+=1
			self.userNbr.setText(str(i))
			self.modSetComboBox.addItem(str(user).strip("(',')"))
			self.delSetComboBox.addItem(str(user).strip("(',')"))
				
	def settingSignal(self) :
		self.logAdd.clicked.connect(self.settingSlot)
		
		self.logModif.clicked.connect(self.settingSlot)
		self.modSetComboBox.currentTextChanged.connect(self.comboSetting)
		
		self.logDel.clicked.connect(self.settingSlot)
		self.delSetComboBox.currentTextChanged.connect(self.comboSetting)
		
		self.settingInfo.clicked.connect(self.settingSlot)
		
		self.setOK.clicked.connect(self.settingSlot)
		self.setAnnuler.clicked.connect(self.settingSlot)
		
	def settingSlot(self):
		if self.sender() == self.logAdd :
			self.settingAdd()
			self.selection()
			
		if self.sender() == self.logModif :
			self.settingModif()
			self.selection()
			
		if self.sender() == self.logDel :
			self.settingDel()
			self.selection()			
		
		if self.sender() == self.settingInfo :
			self.socityInfo()
			self.selection()		
		
		if self.sender() == self.setOK :
			self.selection()	
			self.close()
		
		if self.sender() == self.setAnnuler :	
			self.close()
			
	def comboSetting(self):	
		self.modUserName.setText(self.modSetComboBox.currentText())
		self.delUserName.setText(self.delSetComboBox.currentText())
		
	def settingAdd(self) :
		try:
			query.execute("SELECT user_name FROM Login ORDER BY id ASC")
			list = query.fetchall()
			
			for item in list:	
				if str(item).strip("(',')") == self.newUserName.text() :
					self.newUserName.clear()
					self.messageFactory.raiseMatch(str(item).strip("(',')"))
					
			if self.newUserName.text() != "" and self.newPassWord.text() != "" :
				query.execute("INSERT INTO Login (user_name, user_password) VALUES('"+self.newUserName.text()+"','"+self.newPassWord.text()+"')")
				
				conn.commit()
				self.newUserName.clear()
				self.newPassWord.clear()
				
				self.messageFactory.raiseAdder("Utilisateur")
				self.selection()
				
			else :
				if self.newUserName.text() == "" :
					self.messageFactory.raiseCaseExcept("Nom d'utilisaeur")
				if self.newPassWord.text() == "":
					self.messageFactory.raiseCaseExcept("Mot de passe")
		except:
			self.messageFactory.raiseCharExcept()
		
	def settingModif(self) :
		try:
			if self.modUserName.text() != "" and self.modUserName.text() != "..." and self.modPassWord.text() != ""  \
				 and self.modSetComboBox.currentText() != "..." and self.who.text() == "admin" and self.modUserName.text() != "admin" :
				 
				query.execute("UPDATE Login SET user_name='"+self.modUserName.text()+"', user_password='"+self.modPassWord.text()+"'\
				WHERE user_name = '"+self.modSetComboBox.currentText()+"'")
				conn.commit()
			
				self.messageFactory.raiseModifier("Utilisateur")
				
				self.modUserName.clear()
				self.modPassWord.clear()
				self.selection()
					
			else :
				if self.modUserName.text() == "" :
					self.messageFactory.raiseCaseExcept("Nouveau nom d'utilisateur")
				if self.modPassWord.text() == "" :
					self.messageFactory.raiseCaseExcept("Nouveau mot de passe")
				if self.modUserName.text() == "admin" :
					self.messageFactory.raiseIndefinedExcept("Vous ne pouvez pas modifier le nom d'utlisateur admin")
				if self.user != "admin" :
					self.messageFactory.raiseIndefinedExcept("Seul admin est abilité à modifier ou supprimer un compte")
					
				self.modUserName.clear()
				self.modPassWord.clear()
		except:
			self.messageFactory.raiseCharExcept()
			
	def settingDel(self) :
		# try:
		self.testingco = 0
		if self.delUserName.text() != "" and self.delUserName.text() != "..." and self.delSetComboBox.currentText() != "..." and self.who.text() == "admin":
			
			if self.delUserName.text() != "admin" :
				query.execute("DELETE FROM 'Login' WHERE user_name = '"+self.delSetComboBox.currentText()+"'")
				conn.commit()
			
				self.messageFactory.raiseDeleter("Utilisateur")
			
				self.delUserName.clear()
				self.selection()
				self.testingco = 1
			else :
				self.messageFactory.raiseIndefinedExcept("Vous ne pouvez pas supprimer le nom d'utlisateur admin")
			
		else :
			if self.delUserName.text() == "" :
				self.messageFactory.raiseCaseExcept("le Nom d'utilisateur à supprimer")
			if self.user != "admin" :
				self.messageFactory.raiseIndefinedExcept("Seul admin est abilité à modifier ou supprimer un compte")
			self.delUserName.clear()
		# except:
			# self.messageFactory.raiseCharExcept()
		
	def socityInfo(self):
		try:
			if self.societyName.text() != "" and self.societyNumber.text() != "" :
				query.execute("UPDATE Society SET society_name='"+self.societyName.text()+"', society_number='"+self.societyNumber.text()+"' WHERE id = 1 ")
				conn.commit()
							
				self.messageFactory.raiseModifier("Information")
				
				self.modUserName.clear()
				self.modPassWord.clear()
				
			else :
				self.messageFactory.raiseCaseExcept("toutes les cases")
				if self.societyName.text() == "" :
					self.messageFactory.raiseCaseExcept("le nom de la society")
				elif self.societyNumber.text() == "" :
					self.messageFactory.raiseCaseExcept("le numéro de la society")
		except:
			self.messageFactory.raiseCharExcept()
				
#============================================================================================================CONNECTION 		
	
qtValid= "DESIGN/DIALOGS/connecterDialog.ui"
Ui_ConnectDialog, QtBaseClass = uic.loadUiType(qtValid)

class ConnectDialog(QDialog, Ui_ConnectDialog):#CONFIRM : VALID DIALOG ON VLDButton CLICK()
	def __init__(self):
		QDialog.__init__(self)
		Ui_ConnectDialog.__init__(self)
		self.setupUi(self)
		self.setWindowTitle("Connexion au systeme")
		
		self.MF = MessageFactory()
		self.Set = SettingCreatorDialog()
		#Connecter signals
		self.userName.textChanged.connect(self.Texter)
		self.userPassword.textChanged.connect(self.Texter)
		
		self.disconnect.clicked.connect(self.connSlot)
	
	def Texter(self):
		self.label.setText("")
		
	def connSlot(self):
		if self.sender() == self.disconnect :
			self.close()
			
	def login(self,window) :
		query.execute("SELECT user_name FROM Login ORDER BY id ASC")
		self.users=[]
		self.users = query.fetchall()
		query.execute("SELECT user_password FROM Login ORDER BY id ASC")
		self.passwords=[]
		self.passwords=query.fetchall()
		i=0
		j=0
		while i < len(self.users) :
			while j < len(self.passwords) :
				if self.userName.text() == str(self.users[i]).strip("(',')")  and self.userPassword.text() == str(self.passwords[j]).strip("(',')") :
					query.execute("UPDATE User SET user='"+self.userName.text()+"' WHERE id = 1")
					conn.commit()
					window.show()
					
					self.Set.selection()
					
					self.userName.setText("")
					self.userPassword.setText("")
					self.label.setText("")
					self.close()
				else :
					self.label.setText("Identification impossible, veuillez réessayer.")
				j+=1
				i+=1	

# =============PRODUCTCREATOR DIALOG		
qtProdDialog= "DESIGN/DIALOGS/ProductCreatorDialog.ui"
Ui_ProductCreatorDialog, QtBaseClass = uic.loadUiType(qtProdDialog)	
class ProductCreatorDialog(QDialog, Ui_ProductCreatorDialog):#EDIT : MODIF Product Name,Price DIALOG
	
	def __init__(self):
		QDialog.__init__(self)
		Ui_ProductCreatorDialog.__init__(self)
		self.setupUi(self)
		# self.show()
		
		self.setWindowTitle("Edition des produits.")
	#======INSTANCES
		self.MF = MessageFactory()
		
		self.CategorieInit()
		self.INTEGERS()

	#======SIGNALS PRODUCTS AND CATEGORIES
		self.addCat.clicked.connect(self.AddCat)
		self.delCat.clicked.connect(self.AddCat)
		
		self.addProdBC.clicked.connect(self.AddProd)
		self.modifPBC.clicked.connect(self.ModifProd)
		self.addProd.clicked.connect(self.AddProd)
		self.modifProd.clicked.connect(self.ModifProd)
		self.delProd.clicked.connect(self.DelProd)
		
		self.modfCatName.currentTextChanged.connect(self.ProductInit)
		self.delProdCat.currentTextChanged.connect(self.ProductInit)
		
		self.delCatName.currentTextChanged.connect(self.AddCat)
		
	#======PROTOTYPES AND FUNCTIONS
	def CategorieInit(self):
		# SELECTING  category_name FROM CATEGORIES -->  CATLIST:
		query.execute("SELECT category_name FROM Categories ORDER by id ASC ")#SELECT : Category Name
		self.catList = list()
		self.catList = query.fetchall()
		#Cleaning
		self.addCatName.clear()
		self.modfCatName.clear()
		self.delCatName.clear()
		self.delProdCat.clear()
		
		#INIT CATS ITEMS
		i=0
		self.addCatName.addItem("")
		self.modfCatName.addItem("")
		self.delCatName.addItem("")
		self.delProdCat.addItem("")
		for prod in self.catList :
			self.addCatName.addItem(str(prod).strip("(',')"))
			self.modfCatName.addItem(str(prod).strip("(',')"))
			self.delCatName.addItem(str(prod).strip("(',')"))
			self.delProdCat.addItem(str(prod).strip("(',')"))
			i+=1
	
	def INTEGERS(self):
		
		self.newProdPriceBC.textChanged.connect(self.Regex)
		self.newProdPrice.textChanged.connect(self.Regex)
		self.newProdStockBC.textChanged.connect(self.Regex)
		self.newProdStocMinBC.textChanged.connect(self.Regex)
		self.newProdStock.textChanged.connect(self.Regex)
		self.newProdStockMin.textChanged.connect(self.Regex)
		self.modifProdPriceBC.textChanged.connect(self.Regex)
		self.modifProdStockBC.textChanged.connect(self.Regex)
		self.modifProdStockMinBC.textChanged.connect(self.Regex)
		self.modifProdPrice.textChanged.connect(self.Regex)
		self.modifProdStock.textChanged.connect(self.Regex)
		self.modifProdStockMin.textChanged.connect(self.Regex)
		
	def Regex(self):	
		try:
			sender = self.sender()
			if sender.text() != "" and sender.text() != 0:
				mod = int(sender.text()) % int(sender.text()) 
			pattern = re.compile("^[1-9]")
			status = re.search(pattern, sender.text())
			if sender == self.newProdPriceBC and self.newProdPriceBC.text() != "Prix du produit" \
			or sender == self.newProdPrice and self.newProdPrice.text() != "Prix du produit" \
			or sender == self.newProdStockBC and self.newProdStockBC.text() != "Stock du produit" \
			or sender == self.newProdStocMinBC and self.newProdStocMinBC.text() != "Stock minimum" \
			or sender == self.newProdStock and self.newProdStock.text() != "Stock du produit" \
			or sender == self.newProdStockMin and self.newProdStockMin.text() != "Stock minimum" \
			or sender == self.modifProdPriceBC and self.modifProdPriceBC.text() != "Nouveau prix" \
			or sender == self.modifProdStockBC and self.modifProdStockBC.text() != "Nouveau stock" \
			or sender == self.modifProdStockMinBC and self.modifProdStockMinBC.text() != "Nouveau stock minimum" \
			or sender == self.modifProdPrice and self.modifProdPrice.text() != "Nouveau prix" \
			or sender == self.modifProdStock and self.modifProdStock.text() != "Nouveau stock" \
			or sender == self.modifProdStockMin and self.modifProdStockMin.text() != "Nouveau stock minimum":
				if sender.text() != "" and not status :
					self.onlyInt = QIntValidator()
					sender.setValidator(self.onlyInt)
					sender.setText("")
					self.MF.raiseNoInt()
			elif sender.text() == "0" :
				sender.setText("")
				self.MF.raiseNoInt()
		except :
			sender.setText("")
			self.MF.raiseNoInt()
				
		
	def ProductItemer(self,category,prodCombo):#SET ITEMS INTO PRODCOMBOLIST
			query.execute("SELECT product_name FROM Products WHERE product_category ='"+category.currentText()+"' ORDER by id ASC ")#SELECT : Category Name
			prodList = []
			prodList = query.fetchall()
			
			prodCombo.clear()
			prodCombo.addItem("...")
			i=0
			for prod in prodList :
				prodCombo.addItem(str(prod).strip("(',')"))
				i+=1
	
	def ProductInit(self):
		sender = self.sender()
		if sender == self.modfCatName :
			self.ProductItemer(self.modfCatName, self.prodToModif)
		elif sender == self.delProdCat :
			self.ProductItemer(self.delProdCat, self.delProdName)

	#======VALIDATION SLOTS
	def AddProd(self) :
		sender = self.sender()
		if sender == self.addProdBC :
			a = self.newProdNameBC.text()
			b = self.newProdPriceBC.text()
			c = self.newProdBC.text()
			d = self.newProdStockBC.text()
			e = self.newProdStocMinBC.text()
			try:
				if  self.newProdNameBC.text() != 'Nom du produit' and self.newProdPriceBC.text() != 'Prix du produit'\
				and self.newProdBC.text() != 'Code bar du produit'	and self.newProdStockBC.text() != 'Stock du produit'\
				and self.newProdStocMinBC.text() != 'Stock minimum' and a !="" and b !="" and c !="" and d !="" and e !="":
				
					query.execute('insert into Products \
					(product_name,product_price,product_BC,product_stock,product_stock_min,product_date)\
					values("{0}","{1}" ,"{2}" ,"{3}","{4}","{5}" );'\
					.format(self.newProdNameBC.text(),self.newProdPriceBC.text(),self.newProdBC.text(),self.newProdStockBC.text(),\
					self.newProdStocMinBC.text(),self.newProdDateBC.text()))
					conn.commit()
					
					self.newProdNameBC.clear()
					self.newProdPriceBC.clear()
					self.newProdBC.clear()
					self.newProdStockBC.clear()
					self.newProdStocMinBC.clear()
					
					self.MF.raiseAdder("Produit")
					
					
				else :  
					self.MF.raiseCaseExcept("Toutes les cases")
			except:
				self.MF.raiseCharExcept()
				
		elif sender == self.addProd :
			a = self.newProdName.text()
			b = self.newProdPrice.text()
			c = self.newProdStock.text()
			d = self.newProdStockMin.text()
			try:
				if  self.newProdName.text() != 'Nom du produit' and self.newProdPrice.text() != 'Prix du produit'\
				and self.newProdStock.text() != 'Stock du produit' and self.newProdStockMin.text() != 'Stock minimum'\
				and a !="" and b !="" and c !="" and d !="" :
					query.execute('insert or replace into Products \
					(product_category,product_name,product_price,product_stock,product_stock_min,product_date)\
					values("{0}","{1}" ,"{2}" ,"{3}","{4}","{5}" );'\
					.format(self.addCatName.currentText(),self.newProdName.text(),self.newProdPrice.text(),self.newProdStock.text(),\
					self.newProdStockMin.text(),self.newProdDate.text()))
					conn.commit()
					
					self.newProdName.clear()
					self.newProdPrice.clear()
					self.newProdStock.clear()
					self.newProdStockMin.clear()
					
					self.MF.raiseAdder("Produit")
					
				else :  
					self.MF.raiseCaseExcept("Toutes les cases")
			except:
				self.MF.raiseCharExcept()

	def ModifProd(self) :
		sender = self.sender()
		
		if sender == self.modifPBC :
			try:
				if   self.BCToModif.text() != ''  and self.BCToModif.text() != 'Code bar du produit à modifier'  :
				
					a = self.modifProdNameBC.text() 
					b = self.modifProdPriceBC.text() 
					c = self.modifProdBC.text() 
					d = self.modifProdStockBC.text() 
					e = self.modifProdStockMinBC.text()
					
					f = self.BCToModif.text()
					
					if self.modifProdNameBC.text() != '' and a != "Nouveau nom":
						query.execute('update Products set product_name ="{0}" where product_BC = "{1}"'.format(a,f))
					if self.modifProdPriceBC.text() != '' and b != "Nouveau prix" :
						query.execute('update Products set product_price ="{0}" where product_BC = "{1}"'.format(b,f))
					if self.modifProdBC.text() != '' and c != "Nouveau code bar" :
						query.execute('update Products set product_BC="{0}" where product_BC = "{1}"'.format(c,f))
					if self.modifProdStockBC.text() != '' and d != "Nouveau stock" :
						query.execute('update Products set product_stock="{0}" where product_BC = "{1}"'.format(d,f))
					if self.modifProdStockMinBC.text() != '' and e != "Nouveau stock minimum" :
						query.execute('update Products set product_stock_min ="{0}" where product_BC = "{1}"'.format(e,f))
					
					conn.commit()
					
					self.BCToModif.clear()
					self.modifProdNameBC.clear()
					self.modifProdPriceBC.clear()
					self.modifProdBC.clear()
					self.modifProdStockBC.clear()
					self.modifProdStockMinBC.clear()
					
					self.MF.raiseModifier("Produit")
					
					
				else :  
					self.MF.raiseCaseExcept("Toutes les cases")
			except:
				self.MF.raiseCharExcept()
				
		elif sender == self.modifProd :
			try :
				if   self.modfCatName.currentText() != '' and self.prodToModif.currentText() != '...' :
				
					a = self.modifProdName.text() 
					b = self.modifProdPrice.text() 
					c = self.modifProdStock.text() 
					d = self.modifProdStockMin.text()
					
					e = self.modfCatName.currentText()
					g = self.prodToModif.currentText()
					
					if self.modifProdName.text() != ''  and a != "Nouveau nom":
						query.execute('update Products set product_name ="{0}" where product_category = "{1}" AND product_name="{2}"'.format(a,e,g))
					if self.modifProdPrice.text() != '' and b != "Nouveau prix":
						query.execute('update Products set product_price ="{0}" where product_category = "{1}" AND product_name="{2}"'.format(b,e,g))
					if self.modifProdStock.text() != '' and c != "Nouveau stock":
						query.execute('update Products set product_stock="{0}" where product_category = "{1}" AND product_name="{2}"'.format(c,e,g))
					if self.modifProdStockMin.text() != '' and d != "Nouveau stock minimum":
						query.execute('update Products set product_stock_min="{0}" where product_category = "{1}" AND product_name="{2}"'.format(d,e,g))
					
					conn.commit()
					
					self.modifProdName.clear()
					self.modifProdPrice.clear()
					self.modifProdStock.clear()
					self.modifProdStockMin.clear()
					
					self.prodToModif.clear()
					self.CategorieInit()
					
					self.MF.raiseModifier("Produit")
					
					
				else :  
					self.MF.raiseCaseExcept("Toutes les cases")
			except:
				self.MF.raiseCharExcept()
				
	def DelProd(self) :
		try :
			if self.delProdNameBC.text() != '' and self.delProdNameBC.text() != 'Code bar du produit à supprimer':
				query.execute ('delete from Products where product_BC ="{0}" ;'.format( self.delProdNameBC.text()) )
				conn.commit()
				
				self.delProdNameBC.setText('Code bar du produit à supprimer')
				
				self.MF.raiseDeleter("Produit")
				
			elif self.delProdCat.currentText() != "" and self.delProdCat.currentText() != "Catégorie du produit à supprimer" and self.delProdName.currentText() != "..." :
				query.execute ('delete from Products where product_name ="{0}" ;'.format(self.delProdName.currentText()))
				conn.commit()
				
					
				self.delProdName.clear()
				self.CategorieInit()
				
				self.MF.raiseDeleter("Produit")
					
			else :  
				self.MF.raiseCaseExcept("toutes les cases")
		except:
			self.MF.raiseCharExcept()
			
	def AddCat(self) :
		sender = self.sender()
		
		if sender == self.addCat:
			try:
				if self.addNewCat.text() != "" and self.addNewCat.text() != "Nouvelle catégorie":
					query.execute('insert into Categories (category_name) values ("{0}");'.format(self.addNewCat.text()))
					conn.commit()
					
					self.addNewCat.clear()
						
					self.MF.raiseAdder("Catégorie")
					self.CategorieInit()
						
				else :  
					self.MF.raiseCaseExcept("la case catégorie")
			except:
				self.MF.raiseCharExcept()
				
		if sender == self.delCat:
			try:
				a = self.delCatName.currentText()
				if self.delCatName.currentText() != "" and self.delCatName.currentText() != "Catégorie à supprimer"  :
					# self.MF.raiseIndefinedExcept("Attention on supprimant la catégorie\n tous les produits qu'elle contient le seront aussi")
					query.execute('DELETE FROM Products WHERE product_category = "{0}"'.format(a))
					query.execute ('DELETE FROM Categories WHERE category_name ="{0}" ;'.format(self.delCatName.currentText()))
					conn.commit()
				
					self.MF.raiseDeleter("Catégorie")
					self.CategorieInit()
					
				else :  
					self.MF.raiseCaseExcept("la case catégorie à supprimer")
			except:
				self.MF.raiseCharExcept()
		
		if sender == self.delCatName and sender.currentText() !="":
			self.MF.raiseIndefinedExcept("En cas de suppression tous les produits de la catégorie '"+sender.currentText()+"' seront supprimés")

# =============STOCK DIALOG		
qtStock= "DESIGN/DIALOGS/AlertS.ui"
Ui_AlertS, QtBaseClass = uic.loadUiType(qtStock)
class AlertS (QDialog,Ui_AlertS):
	def __init__ (self,parent=None):
		QDialog.__init__(self)
		Ui_AlertS.__init__(self)
		self.setupUi(self)
		self.setWindowTitle("Alerts Stock.")
		
		self.clear.clicked.connect(self.Clear)
		
		self.InitAlert()
	def Clear(self):
		self.textEdit.clear()
		
	def InitAlert(self) :
		self.textEdit.clear()
		query.execute('select product_name, product_stock,product_stock_min from Products ORDER BY id ASC')
		a = query.fetchall()
		for x in a:
			if x[1] != '' and x[2] != '' and x[1] <= x[2]:
				self.string = 'Le stock de {0} est : {1} (stock minimum {2}) \n'.format(x[0],x[1],x[2])
				self.textEdit.append(self.string)
		
# =============DATE PEREMPTION DIALOG		
qtPDate= "DESIGN/DIALOGS/AlertS.ui"
Ui_AlertP, QtBaseClass = uic.loadUiType(qtPDate)	
class AlertP (QDialog,Ui_AlertP):
	def __init__ (self,parent=None):
		QDialog.__init__(self)
		Ui_AlertP.__init__(self)
		self.setupUi(self)
		
		self.setWindowTitle("Alerts Peremption.")
		
		self.clear.clicked.connect(self.Clear)
		
		self.InitAlert()
		
	def Clear(self):
		self.textEdit.clear()
		
	def InitAlert(self) :
		self.textEdit.clear()
		query.execute('select product_name,product_date from Products ORDER BY id ASC ;')
		fetch = query.fetchall()
		today=datetime.datetime.today().date()
		for x in fetch:
			date_1 = datetime.datetime.strptime(x[1], "%d/%m/%Y").date()
			end_date = date_1 + datetime.timedelta(days=-20)
			if today  >= end_date and today < date_1:
				self.string = 'Il reste moin de 20 jours avant la date de peremption de {0} : {1}\n\n'.format(x[0],x[1])
				self.textEdit.append(self.string)

# =============PRODUCT DIALOG		
qtMyProdDialog= "DESIGN/DIALOGS/MyProduct.ui"
Ui_MyProduct, QtBaseClass = uic.loadUiType(qtMyProdDialog)				
class MyProductDialog (QDialog,Ui_MyProduct):
	def __init__ (self,parent=None):
		QDialog.__init__(self)
		Ui_MyProduct.__init__(self)
		self.setupUi(self)
						
		self.localDateTime = time.strftime("%d-%m-%Y      %H:%M:%S")
		self.localDate = time.strftime("%d-%m-%Y")
		self.Date = time.strftime("%Y-%m-%d")
		
		self.MF = MessageFactory()
		
		self.MyProdInit()
		self.setWindowTitle("Liste des produits.")
		
		self.find.clicked.connect(self.ProdFinder)
		self.findCB.clicked.connect(self.ProdFinderCB)
		self.myProd.clicked.connect(self.MyProdInit)
		
	def MyProdInit(self):
		self.cur = query
        
		conn.text_factory = str
		self.cur.execute(" SELECT product_name , product_price , product_BC, product_category, product_stock, product_stock_min, product_date  FROM Products\
		ORDER BY id ASC")

		self.tableWidget.setColumnCount(7) 
		self.tableWidget.setColumnWidth(0, 120)
		self.tableWidget.setColumnWidth(1, 115)
		self.tableWidget.setColumnWidth(2, 120)
		self.tableWidget.setColumnWidth(3, 115)
		self.tableWidget.setColumnWidth(4, 104)
		self.tableWidget.setColumnWidth(5, 104)
		self.tableWidget.setColumnWidth(6, 115)
		
		self.tableWidget.setHorizontalHeaderLabels(['Produit', 'Prix', 'Codebar','Catégorie', 'Stock', 'Stock min','Date de péremption'])
		self.header = self.tableWidget.horizontalHeader()
		self.header.setDefaultAlignment(Qt.AlignHCenter)
		
		self.tableWidget.setRowCount(0)
		for row, form in enumerate(self.cur):
			self.tableWidget.insertRow(row)
			for column, item in enumerate(form):
				#print(str(item))
				self.tableWidget.setItem(row, column,QTableWidgetItem(str(item)))

	def ProdFinder(self):
		a = self.prodToFind
		try:
			if a.text() != '' and a.text() != 'Rechercher un produit' :
				
				conn.text_factory = str
				query.execute(" SELECT product_name , product_price , product_BC, product_category, product_stock, product_stock_min, product_date  FROM Products\
				WHERE product_name = '{0}' ORDER BY product_date ASC".format(a.text()))
				product = query.fetchall()
				if product != [] :
					self.tableWidget.clearContents()
					self.tableWidget.clear()
					
					self.tableWidget.setColumnCount(7) 
					self.tableWidget.setColumnWidth(0, 120)
					self.tableWidget.setColumnWidth(1, 115)
					self.tableWidget.setColumnWidth(2, 120)
					self.tableWidget.setColumnWidth(3, 115)
					self.tableWidget.setColumnWidth(4, 104)
					self.tableWidget.setColumnWidth(5, 104)
					self.tableWidget.setColumnWidth(6, 115)
					
					self.tableWidget.setHorizontalHeaderLabels(['Produit', 'Prix', 'Codebar','Catégorie', 'Stock', 'Stock min','Date de péremption'])
					self.header = self.tableWidget.horizontalHeader()
					self.header.setDefaultAlignment(Qt.AlignHCenter)
					
					produit=product[0][0]
					prix=product[0][1]
					CB= product[0][2]
					cat= product[0][3]
					stock=product[0][4]
					stockMin=product[0][5]
					date=product[0][6]
					
					#self.tableWidget.setSortingEnabled(False)
					#self.tableWidget.setSortingEnabled(True)
					
					self.tableWidget.setRowCount(0)	
					self.rowPosition = self.tableWidget.rowCount()
					self.tableWidget.insertRow(self.rowPosition)
					self.tableWidget.setItem(self.rowPosition , 0, QTableWidgetItem(str(produit)))
					self.tableWidget.setItem(self.rowPosition , 1, QTableWidgetItem(str(prix)))
					self.tableWidget.setItem(self.rowPosition , 2, QTableWidgetItem(str(CB)))
					self.tableWidget.setItem(self.rowPosition , 3, QTableWidgetItem(str(cat)))
					self.tableWidget.setItem(self.rowPosition , 4, QTableWidgetItem(str(stock)))
					self.tableWidget.setItem(self.rowPosition , 5, QTableWidgetItem(str(stockMin)))
					self.tableWidget.setItem(self.rowPosition , 6, QTableWidgetItem(str(cat)))

					a.setText("Rechercher un produit")
					
				elif product == [] :
					self.MF.raiseNoMatch(str(a.text()))
					a.setText("Rechercher un produit")
				
			else :
				self.MF.raiseCaseExcept("recherche")
				a.setText("Rechercher un produit")
		except:
			self.MF.raiseIndefinedExcept("Une erreur est survenu\nVérifiez l'orthographe du produit")
									
	def ProdFinderCB(self):
		a = self.prodToFindCB
		try:
			if a.text() != '' and a.text() != 'Rechercher un produit avec codebar' :
				
				conn.text_factory = str
				query.execute(" SELECT product_name , product_price , product_BC, product_category, product_stock, product_stock_min, product_date  FROM Products\
				WHERE product_BC = '{0}' ORDER BY product_date ASC".format(a.text()))
				product = query.fetchall()
				if product != [] :
					self.tableWidget.clearContents()
					self.tableWidget.clear()
					
					self.tableWidget.setColumnCount(7) 
					self.tableWidget.setColumnWidth(0, 120)
					self.tableWidget.setColumnWidth(1, 115)
					self.tableWidget.setColumnWidth(2, 120)
					self.tableWidget.setColumnWidth(3, 115)
					self.tableWidget.setColumnWidth(4, 104)
					self.tableWidget.setColumnWidth(5, 104)
					self.tableWidget.setColumnWidth(6, 115)
					
					self.tableWidget.setHorizontalHeaderLabels(['Produit', 'Prix', 'Codebar','Catégorie', 'Stock', 'Stock min','Date de péremption'])
					self.header = self.tableWidget.horizontalHeader()
					self.header.setDefaultAlignment(Qt.AlignHCenter)
					
					produit=product[0][0]
					prix=product[0][1]
					CB= product[0][2]
					cat= product[0][3]
					stock=product[0][4]
					stockMin=product[0][5]
					date=product[0][6]
					
					#self.tableWidget.setSortingEnabled(False)
					#self.tableWidget.setSortingEnabled(True)
					
					self.tableWidget.setRowCount(0)	
					self.rowPosition = self.tableWidget.rowCount()
					self.tableWidget.insertRow(self.rowPosition)
					self.tableWidget.setItem(self.rowPosition , 0, QTableWidgetItem(str(produit)))
					self.tableWidget.setItem(self.rowPosition , 1, QTableWidgetItem(str(prix)))
					self.tableWidget.setItem(self.rowPosition , 2, QTableWidgetItem(str(CB)))
					self.tableWidget.setItem(self.rowPosition , 3, QTableWidgetItem(str(cat)))
					self.tableWidget.setItem(self.rowPosition , 4, QTableWidgetItem(str(stock)))
					self.tableWidget.setItem(self.rowPosition , 5, QTableWidgetItem(str(stockMin)))
					self.tableWidget.setItem(self.rowPosition , 6, QTableWidgetItem(str(cat)))
					
					a.setText("Rechercher un produit avec codebar")
				
				elif product == [] :
					self.MF.raiseNoMatch(str(a.text()))
					a.setText("Rechercher un produit")
				
			else :
				self.MF.raiseCaseExcept("recherche avec codebar")
				a.setText("Rechercher un produit")
									
		except:
			self.MF.raiseIndefinedExcept("Une erreur est survenu\nVérifiez l'orthographe du produit")
				
# =============RECEIPT DIALOG		
qtReceipt= "DESIGN/DIALOGS/Receipt.ui"
Ui_Receipt, QtBaseClass = uic.loadUiType(qtReceipt)		
class Receipt(QDialog,Ui_Receipt):
	def __init__ (self,parent=None):
		QDialog.__init__(self)
		Ui_Receipt.__init__(self)
		self.setupUi(self)
		
		self.setWindowTitle("Recette du jour.")
		self.Date = time.strftime("%Y-%m-%d")
		self.year = time.strftime("%Y")
		self.month = time.strftime("%m")
		self.day = time.strftime("%d")
		
		self.MF = MessageFactory()
		self.MyProdInit()
		
		self.selectDate.clicked.connect(self.DateFinder)
		self.today.clicked.connect(self.MyProdInit)
		self.discount.clicked.connect(self.MyDiscInit)
		
	def MyProdInit(self):
		self.dater.setDisplayFormat("yyyy-MM-dd")
		self.dater.setDate(QtCore.QDate(int(self.year),int(self.month), int(self.day)))
		
		self.cur = query
        
		conn.text_factory = str
		self.cur.execute(' SELECT product_name, product_price, product_qnt,product_date,product_ref FROM Receipt WHERE product_date = "{0}" ORDER BY product_ref ASC'\
		.format(str(self.Date)))

		self.tableWidget.setColumnCount(5) 
		self.tableWidget.setColumnWidth(0, 160)
		self.tableWidget.setColumnWidth(1, 160)
		self.tableWidget.setColumnWidth(2, 160)
		self.tableWidget.setColumnWidth(3, 160)
		self.tableWidget.setColumnWidth(4, 160)
		
		self.tableWidget.setHorizontalHeaderLabels(['PRODUIT', 'PRIX','QUANTITE', 'DATE', 'N°TICKET'])
		self.header = self.tableWidget.horizontalHeader()
		self.header.setDefaultAlignment(Qt.AlignHCenter)
		
		self.tableWidget.setRowCount(0)
		for row, form in enumerate(self.cur):
			self.tableWidget.insertRow(row)
			for column, item in enumerate(form):
				#print(str(item))
				self.tableWidget.setItem(row, column,QTableWidgetItem(str(item)))

	def MyDiscInit(self):
		self.dater.setDisplayFormat("yyyy-MM-dd")
		
		self.cur = query
        
		conn.text_factory = str
		self.cur.execute(' SELECT disc_ref, disc_value, disc_date FROM Discount WHERE disc_date = "{0}" ORDER BY id ASC'\
		.format(str(self.dater.text())))

		self.tableWidget.clearContents() 
		self.tableWidget.setColumnCount(3) 
		self.tableWidget.setColumnWidth(0, 260)
		self.tableWidget.setColumnWidth(1, 260)
		self.tableWidget.setColumnWidth(2, 260)
		
		self.tableWidget.setHorizontalHeaderLabels(['N°TICKET', 'MONTANT REMISE','DATE DE REMISE'])
		self.header = self.tableWidget.horizontalHeader()
		self.header.setDefaultAlignment(Qt.AlignHCenter)
		
		self.tableWidget.setRowCount(0)
		for row, form in enumerate(self.cur):
			self.tableWidget.insertRow(row)
			for column, item in enumerate(form):
				#print(str(item))
				self.tableWidget.setItem(row, column,QTableWidgetItem(str(item)))

	def DateFinder(self):
		a = self.dater
		try:
			if a.text() != '' :
				
				conn.text_factory = str
				self.cur.execute(' SELECT product_name, product_price, product_qnt,product_date,product_ref FROM Receipt WHERE product_date = "{0}" ORDER BY product_ref ASC'\
				.format(str(a.text())))

				self.tableWidget.setColumnCount(5) 
				self.tableWidget.setColumnWidth(0, 160)
				self.tableWidget.setColumnWidth(1, 160)
				self.tableWidget.setColumnWidth(2, 160)
				self.tableWidget.setColumnWidth(3, 160)
				self.tableWidget.setColumnWidth(4, 160)
				
				self.tableWidget.setHorizontalHeaderLabels(['PRODUIT', 'PRIX','QUANTITE', 'DATE', 'N°TICKET'])
				self.header = self.tableWidget.horizontalHeader()
				self.header.setDefaultAlignment(Qt.AlignHCenter)
				
				self.tableWidget.setRowCount(0)
				for row, form in enumerate(self.cur):
					self.tableWidget.insertRow(row)
					for column, item in enumerate(form):
						#print(str(item))
						self.tableWidget.setItem(row, column,QTableWidgetItem(str(item)))
				
			else :
				self.MF.raiseCaseExcept("recherche")
		except:
			self.MF.raiseIndefinedExcept("Une erreur est survenu\nVérifiez si la date est au format YYYY/MM/DD")
									
#=============REGISTER DIALOG
qtRegisterDialog= "DESIGN/DIALOGS/registerCreatorDialog.ui"
Ui_registerCreatorDialog, QtBaseClass = uic.loadUiType(qtRegisterDialog)
class RegisterCreatorDialog(QDialog, Ui_registerCreatorDialog):# EDIT : MODIF Product Catégory DIALOG

	def __init__(self):
		QDialog.__init__(self)
		Ui_registerCreatorDialog.__init__(self)
		self.setupUi(self)

		self.setWindowTitle("Mouvements de la Caisse et Des Sommes d'Argent")
		self.referenceSelector = ReferenceSelector()
		self.referenceSelector.refSelector()
		self.referenceSelector.societySelctor()
		
		#===========LOGICALS
		self.Date = time.strftime("%Y-%m-%d")
		self.localDateTime = time.strftime("%d-%m-%Y      %H:%M:%S")
		self.localDate = time.strftime("%d-%m-%Y")
		
		#==========INIT INSTANCES
		self.productDialog = ProductCreatorDialog()
		# self.registerInit = InitDialog()
		self.messageFactory = MessageFactory()
		
		self.RegisterInit()
		self.registerSignal()
		self.registerSum()#somme actuel en caisse
		self.INTEGERS()#somme actuel en caisse
		self.regTotal = []#sum in register
		
		#===========METHODES
		
	def INTEGERS(self):
		
		self.sumInit.textChanged.connect(self.Regex)
		self.depVal.textChanged.connect(self.Regex)
		self.addVal.textChanged.connect(self.Regex)
		
	def Regex(self):
		try:
			sender = self.sender()
			if sender.text() != "" and sender.text() != '0':
				mod = int(sender.text()) % int(sender.text()) 
			# if not re.match("^[0-9]", self.newProdPriceBC.text()) and self.newProdPriceBC.text() != "":
			pattern = re.compile("^[1-9]")
			status = re.search(pattern, sender.text())
			if not status and sender.text() != "":
				self.onlyInt = QIntValidator()
				sender.setValidator(self.onlyInt)
				sender.setText("")
				self.messageFactory.raiseNoInt()
			elif sender.text() == "0" :
				sender.setText("")
				self.messageFactory.raiseNoInt()
		except:
			sender.setText("")
			self.messageFactory.raiseNoInt()

	def registerSum(self):#Selection of ADDS & Depots, binding REGISTER DATA, SETTEXT TO LABELS
		try:
			self.regTotal = []#sum in register
			#SELECTION FROM Register
			query.execute("SELECT register_recette_total FROM Register WHERE register_date = '"+self.Date+"'")
			self.recette = str(query.fetchone()).strip("(',')")	
			self.regTotal.append(int(self.recette))
			
			query.execute("SELECT register_sum_init FROM Register ORDER BY id DESC")
			self.suminit = str(query.fetchone()).strip("(',')")
			self.regTotal.append(int(self.suminit))
			
			query.execute("SELECT register_depense_total FROM Register WHERE register_date = '"+self.Date+"'")
			self.depense = str(query.fetchone()).strip("(',')")
			
			#BINDING DATA ON REGISTER (DEP,ADD)
			self.dep = []
			query.execute("SELECT dep_value FROM Register_dep WHERE dep_date = '"+self.Date+"'")
			self.depensesList = []
			self.depensesList = query.fetchall()
			i=0
			for dep in self.depensesList :	
				self.dep.append(str(dep).strip("(',')"))
				self.dep[i] = int(self.dep[i])
				i+=1
			query.execute("UPDATE Register SET register_depense_total = "+str(sum(self.dep))+" WHERE register_date ='"+self.Date+"'")
		
			self.add = []
			query.execute("SELECT add_value FROM Register_add WHERE add_date = '"+self.Date+"'")
			self.addList = []
			self.addList = query.fetchall()
			i=0
			for add in self.addList :	
				self.add.append(str(add).strip("(',')"))
				self.add[i] = int(self.add[i])
				i+=1
			query.execute("UPDATE Register SET register_ajout_total = "+str(sum(self.add))+" WHERE register_date ='"+self.Date+"'")
			
			conn.commit()
			
			query.execute("SELECT register_ajout_total FROM Register WHERE register_date = '"+self.Date+"'")
			self.ajout = str(query.fetchone()).strip("(',')")
			self.regTotal.append(int(self.ajout))
			
			self.sumInit.setText(self.suminit)
			self.sumDay.setText(str(self.recette))
			self.sumDep.setText(self.depense)
			self.sumAdd.setText(str(self.ajout))
			
			self.sumTotal.display(sum(self.regTotal) - int(self.depense))
			self.actual = str(sum(self.regTotal) - int(self.depense))
			self.actual = str(self.actual)	
			
		except:
			self.messageFactory.raiseIndefinedExcept("au niveau de la caisse")
		
	def RegisterInit(self) :#SELECT AND TEST DAY REGSTER DATA to bind
		
		query.execute("SELECT register_date FROM Register ORDER BY id DESC")
		self.lastDate = query.fetchall()
		
		if self.lastDate == [] :
			query.execute("INSERT INTO Register (register_date, register_sum_init, register_recette_total,register_depense_total,register_ajout_total) VALUES ('"\
			+self.Date+"',0,0,0,0)")
			self.recette = 0
			
		elif str(self.lastDate[0]).strip("(',')") != self.Date :
			query.execute("INSERT INTO Register (register_date, register_sum_init, register_recette_total,register_depense_total,register_ajout_total) VALUES ('"\
			+self.Date+"',0,0,0,0)")
			self.recette = 0
		
		conn.commit()

	def registerSignal(self):
	
		self.depAdd.clicked.connect(self.registerSlot)
		self.addAdd.clicked.connect(self.registerSlot)
		
		self.registerPrint.clicked.connect(self.registerSlot)
		self.editInit.clicked.connect(self.registerSlot)
		
	def registerSlot(self) :	
		self.regTotal = []
		
		if self.sender() == self.depAdd :
			self.newDepense()
			
		if self.sender() == self.editInit and self.sumInit.text() != "":
			query.execute("UPDATE Register SET register_sum_init ={0} WHERE register_date ='{1}'".format(int(self.sumInit.text()),self.Date))
			self.registerSum()
			if self.sumInit.text() != "" :
				self.messageFactory.raiseNoInt()
				
		if self.sender() == self.addAdd :
			self.newDepot()
		
		if self.sender() == self.registerPrint :
			self.registerSum()
			self.printRegister()

	def newDepot(self):
		try:
			if self.addVal.text() != '0' and self.addName.text() != "" :
				query.execute("INSERT INTO Register_add (add_date, add_depo, add_description,add_value) VALUES ('"\
				+self.Date+"','"+str(self.addName.text())+"','"+str(self.addDesc.text())+"','"+str(self.addVal.text())+"')")
				conn.commit()
				 
				
				self.messageFactory.raiseAdder("Dépot")
		
				self.addName.clear()
				self.addDesc.clear()
				self.addVal.clear()
				self.registerSum()
				
			else : 
				self.messageFactory.raiseCaseExcept("Toutes les cases")
				if  self.addName.text() == '':
					self.messageFactory.raiseCaseExcept("la case Dépositaire")
				if  self.addVal.text() == '0':
					self.messageFactory.raiseCaseExcept("la case Somme à dépenser")
		except:
			self.messageFactory.raiseCharExcept()
			
	def newDepense(self):
		try:
			if self.depVal.text() != '0' and self.depDesc.text() != "" :
				query.execute("INSERT INTO Register_dep (dep_date, dep_type, dep_description,dep_value) VALUES ('"\
				+self.Date+"','"+self.depTypeCmbBox.currentText()+"','"+self.depDesc.text()+"','"+self.depVal.text()+"')")			
				conn.commit()
				
				self.messageFactory.raiseAdder("Dépense")
		
				self.depDesc.clear()
				self.depVal.clear()
				self.registerSum()
			else : 
				self.messageFactory.raiseCaseExcept("Toutes les cases")
				if  self.depDesc.text() == '':
					self.messageFactory.raiseCaseExcept("la case déscription")
				if  self.depVal.text() == '0':
					self.messageFactory.raiseCaseExcept("la case Somme à déposer")
		except:
			self.messageFactory.raiseCharExcept()

	def	printRegister(self) :
		try :			
			TICKET = Document()
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")

			h = TICKET.add_heading("                        "+self.referenceSelector.society+" : 'Mouvement de Caisse'", level=1)
			h.bold = True
			h.italic = True
			p = TICKET.add_paragraph("				             "+self.Date)
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_heading("                                     Somme Actuelle en Caisse : "+str(self.actual)+" DA ", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			#DEPENSES
			p = TICKET.add_heading("DEPENSES : ", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			query.execute("SELECT dep_type FROM Register_dep WHERE dep_date = '"+self.Date+"' ORDER BY id")
			self.typeList = []
			self.typeList = query.fetchall()
			query.execute("SELECT dep_description FROM Register_dep WHERE dep_date = '"+self.Date+"' ORDER BY id")
			self.descList = []
			self.descList = query.fetchall()
			query.execute("SELECT dep_value FROM Register_dep WHERE dep_date = '"+self.Date+"' ORDER BY id")
			self.valList = []
			self.valList = query.fetchall()
			
			pi = 0 
			tab = TICKET.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = 'Type'
			heading_cells[1].text = 'Description'
			heading_cells[2].text = 'Valeur'
			heading_cells[3].text = 'Date'
			pi =0
			for dep in self.typeList :
				cells = tab.add_row().cells
				cells[0].text = str(self.typeList[pi]).strip("(',')")
				cells[1].text = str(self.descList[pi]).strip("(',')")
				cells[2].text = str(self.valList[pi]).strip("(',')") + " DA"
				cells[3].text = self.Date
				pi+=1
				
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			#AJOUTS		
			p = TICKET.add_heading("DEPOTS : ", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			query.execute("SELECT add_depo FROM Register_add WHERE add_date = '"+self.Date+"' ORDER BY id")
			self.typeList = []
			self.typeList = query.fetchall()
			query.execute("SELECT add_description FROM Register_add WHERE add_date = '"+self.Date+"' ORDER BY id")
			self.descList = []
			self.descList = query.fetchall()
			query.execute("SELECT add_value FROM Register_add WHERE add_date = '"+self.Date+"' ORDER BY id")
			self.valList = []
			self.valList = query.fetchall()
			
			pi = 0 
			tab = TICKET.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = 'Depositaire'
			heading_cells[1].text = 'Description'
			heading_cells[2].text = 'Valeur'
			heading_cells[3].text = 'Date'
			pi =0
			for dep in self.typeList :
				cells = tab.add_row().cells
				cells[0].text = str(self.typeList[pi]).strip("(',')")
				cells[1].text = str(self.descList[pi]).strip("(',')")
				cells[2].text = str(self.valList[pi]).strip("(',')") + " DA"
				cells[3].text = self.Date
				pi+=1
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_paragraph("SOMME INITIALE EN CAISSE : " +self.suminit+ " DA")
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_paragraph("RECETTE DU JOUR  : " +self.recette+ " DA")
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_heading("\t\tReference No : "+str(self.referenceSelector.ref)+". --"+self.referenceSelector.society+"-- "+self.localDateTime, level=3)
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			TICKET.save('DOCUMENTS/CAISSE/Caisse'+str(self.referenceSelector.ref)+"_"+self.localDate+'.docx' )
			os.startfile('DOCUMENTS\CAISSE\Caisse'+str(self.referenceSelector.ref)+"_"+self.localDate+'.docx' , 'print')
			
			conn.commit()
			
		except :
			self.messageFactory.raisePrintExcept("Caisse"+str(self.ref)+"_"+self.localDate)
	
#==========MAIN CLASS
#========== PRODUCT WIDGET UIC CONVERT & Load
qtProductWidget = "DESIGN/WIDGETS/ProductWidget.ui" # Enter file here.
Ui_ProductWidget, QtBaseClass = uic.loadUiType(qtProductWidget)
class ProductWidget(QWidget, Ui_ProductWidget):
	def __init__(self):
		QWidget.__init__(self)
		Ui_ProductWidget.__init__(self)
		self.setupUi(self)
		# self.show()
		self.setWindowTitle("My Manager Solution.")
		# self.setWindowFlags(Qt.WindowStaysOnTopHint)
						
		self.localDateTime = time.strftime("%d-%m-%Y      %H:%M:%S")
		self.localDate = time.strftime("%d-%m-%Y")
		self.Date = time.strftime("%Y-%m-%d")
		
		#======INSTANCES
		self.MF = MessageFactory()
		self.prodEdit = ProductCreatorDialog()
		self.myProd = MyProductDialog()
		self.alertSt = AlertS()
		self.alertPr = AlertP()
		self.receipt = Receipt()
		self.register = RegisterCreatorDialog()
		
		self.referenceSelector = ReferenceSelector()
		self.society = self.referenceSelector.society 
		self.numero = self.referenceSelector.numero  
		
		#======INIT FUNCTIONS
		self.TableWidgetInit()
		self.Reference()
		self.ButtonSignal()
		self.CatInit()
		
		#======DEFINE
		self.a_list = []#CONTAINER OF SUM TOTAL
		self.produit = 0
		self.prodQnt.setText("")
		#PROD LISTS
		self.pName = []#container List : selected products name
		self.pPrice = []#container List : selected products price
		self.pQnt = []#container List : selected products quantity
		self.pStock = []#container List : selected products stock
		self.pTotalList = []#container List : selected products quantity  
		
		#PROD LISTS to Receipt funct
		self.RpName = []#container List : selected products name
		self.RpPrice = []#container List : selected products price
		self.RpQnt = []#container List : selected products quantity
		self.RpTotalList = []#container List : selected products quantity  
		
		self.INTEGERS()#somme actuel en caisse
		
#===========METHODES		
	def INTEGERS(self):
		
		self.prodPrice.textChanged.connect(self.Regex)
		self.prodQnt.textChanged.connect(self.Regex)
		self.received.textChanged.connect(self.Regex)
		self.discount.textChanged.connect(self.Regex)
		
	def Regex(self):	
		try:
			sender = self.sender()
			if sender.text() != "" and sender.text() != '0':
				mod = int(sender.text()) % int(sender.text()) 
			# if not re.match("^[0-9]", self.newProdPriceBC.text()) and self.newProdPriceBC.text() != "":
			pattern = re.compile("^[0-9 \-]+$")
			status = re.search(pattern, sender.text())
			if not status and sender.text() != "" :
			# To allow only int
				self.onlyInt = QIntValidator()
				sender.setValidator(self.onlyInt)
				sender.setText("")
				self.MF.raiseNoInt()
			elif sender.text() == "0" :
				sender.setText("")
				self.MF.raiseNoInt()
				if sender == self.received and sender.text() == "" or sender.text() == "0" :
					self.payBack.clear()
		except:
			sender.setText("")
			self.MF.raiseNoInt()
			if sender == self.received :
				self.payBack.clear()
		
		if sender == self.received and sender.text() == "" or sender.text() == "0" :
			self.payBack.clear()
		
	def TableWidgetInit(self):	#INIT TABLE WIDGET COLUMN, ALIGNEMENT
		self.tableWidget.setColumnCount(7)
		self.tableWidget.setColumnWidth(0, 145)
		self.tableWidget.setColumnWidth(1, 125)
		self.tableWidget.setColumnWidth(2, 105)
		self.tableWidget.setColumnWidth(3, 125)
		self.tableWidget.setColumnWidth(4, 145)
		self.tableWidget.setColumnWidth(5, 105)
		self.tableWidget.setColumnWidth(6, 125)
		
		self.tableWidget.setHorizontalHeaderLabels(['Produit', 'Prix', 'Quantite','Total', 'Codebar', 'Stock','Date de péremption'])
		self.header = self.tableWidget.horizontalHeader()
		self.header.setDefaultAlignment(Qt.AlignHCenter)
		
	def Reference(self):#TICKET REF SELECTION
		query.execute ('select ref from Ref where id = 1')
		self.ref = query.fetchone()
		self.ref = int(str(self.ref).strip("(',')"))
		self.refNum.setText(str(self.ref))

	def CatInit(self):#SELECTION AND NIT OF CATEGORIES, set SIGNAL TO ButtonInit()
		self.myCats = [self.cat_1, self.cat_2, self.cat_3, self.cat_4, self.cat_5, self.cat_6, self.cat_7, self.cat_8, self.cat_9, self.cat_10]
		for cat in self.myCats :
			cat.setText("")
		
		query.execute('select category_name FROM Categories ORDER BY id ASC')
		self.catList = query.fetchall()
		icon = "IMAGES/Icons/menu.png"
		i=0
		for cat in self.catList:
			self.myCats[i].setText(str(cat).strip("(',')"))
			self.myCats[i].setIcon(QIcon(icon))
			i+=1		
		for cat in self.myCats :
			cat.clicked.connect(self.ButtonInit)
			
#==========BUTTONS UTILITIES	
	def ButtonSignal(self):#SIGNALS OF BUTTONS
		
		self.myProducts.clicked.connect(self.MyProds)
		self.editProduct.clicked.connect(self.ProdEdit)
		
		self.alertS.clicked.connect(self.alerteStockFunc)
		self.alertP.clicked.connect(self.alertePeremption)
		
		self.reiceptButton.clicked.connect(self.ReiceptFunc)
		self.registerButton.clicked.connect(self.RegisterFunc)
		
		#VALIDATION FROM PROD CREATOR DIALOG
		self.prodEdit.addCat.clicked.connect(self.CatInit)
		self.prodEdit.delCat.clicked.connect(self.CatInit)
		
		#VALIDATION BUTTONS SIGNALS
		self.validButton.clicked.connect(self.ValidFunction)
		self.addProdBC.clicked.connect(self.ButtonBC)
		self.addProd.clicked.connect(self.OtherProdFunction)
		self.undoButton.clicked.connect(self.UndoFunction)
		self.cancelButton.clicked.connect(self.CancelFunction)
		
		#CHANGE CALCULATION SIGNAL
		self.received.textChanged.connect(self.ChangeSlot)
		self.discount.textChanged.connect(self.ChangeSlot)
		
		self.myButton = [self.prod_1, self.prod_2, self.prod_3, self.prod_4, self.prod_5, self.prod_6, self.prod_7, self.prod_8, self.prod_9, self.prod_10, \
						self.prod_11, self.prod_12, self.prod_13, self.prod_14, self.prod_15, self.prod_16, self.prod_17, self.prod_18, self.prod_19, self.prod_20]
 
		for button in self.myButton :
			button.clicked.connect(self.ButtonSlot)
	
	def ButtonInit(self):#SLOT AND SELECTION PRODUCTS, INIT STYLESHEET OF BUTTONS
		sender = self.sender()
		if sender.text() != "":
			self.CatIndicator.setText(sender.text())
			#CLEAN 
			c=0
			for button in self.myButton :
				self.myButton[c].setText("")
				self.myButton[c].setStyleSheet("QPushButton{background-color: #f9fafc;text-align:left; color : #000028 ; border: 0px solid #000028 }\
					\n QPushButton:hover{color:white;background-color:#f9fafc; text-decoration:underline;}\
					QPushButton:focus{background-color:#f9fafc;color:#f96332; text-decoration:underline;}")
				c+=1
				
			query.execute("SELECT product_name FROM Products WHERE product_category = '"+str(sender.text())+"' ORDER BY product_name ASC")
			products = query.fetchall()
			
			#INIT
			i=0
			for prod in products :
				self.myButton[i].setText(str(prod).strip("(',')"))
				self.myButton[i].setStyleSheet("QPushButton{background-color:rgb(194,235,237); color :  #000028; border: 0px solid #2f4f4f  ; border-radius: 0.5px; height:30px;}\
					\n QPushButton:hover{color:#f96332;background-color:#ecf0f5;border: 0px solid #2f4f4f }\
					\nQPushButton:focus{color:#f96332;background-color:white; border: 0px solid #2f4f4f ;}")
				i+=1
		else :
			self.CatIndicator.setText("Produits")
			c=0
			for button in self.myButton :
				self.myButton[c].setText("")
				self.myButton[c].setStyleSheet("QPushButton{background-color: #f9fafc;text-align:left; color : #000028 ; border: 0px solid #000028 }\
					\n QPushButton:hover{color:white;background-color:#f9fafc; text-decoration:underline;}\
					QPushButton:focus{background-color:#f9fafc;color:#f96332; text-decoration:underline;}")
				c+=1

	def ButtonSlot(self):#SLOT : BUTTON CLICKED, SELECTION, SETITEM TO TABLEWIDGET
		name = str(self.sender().text())
		if name != '':
			
			self.produit = self.produit + 1
			self.viewProdNum.setText(str(self.produit))
			query.execute('SELECT product_price,product_stock,product_date,product_category FROM Products WHERE product_name = "{0}"'.format(str(name)))
			fetch=query.fetchall()
			
			produit=name
			self.prix=fetch[0][0]
			if self.prodQnt.text()!="":
				self.quantite = int(self.prodQnt.text())
			else:
				self.quantite= 1
                
			self.total= self.prix * self.quantite
			self.stock=fetch[0][1]
			date=fetch[0][2]
			cat=fetch[0][3]
			
			#APPENDING ON CONTAINERS LISTS
			self.pName.append(produit)
			self.pPrice.append(self.prix)
			self.pQnt.append(self.quantite)
			self.pStock.append(int(self.stock))
			self.pTotalList.append(int(self.quantite) * int(self.prix))
			
			self.a_list.append(self.total)
			self.viewTotal.display(sum(self.a_list))
			
			#self.tableWidget.setSortingEnabled(False)
			self.rowPosition = self.tableWidget.rowCount()
			self.tableWidget.insertRow(self.rowPosition)
			#self.tableWidget.setSortingEnabled(True)
			self.tableWidget.setItem(self.rowPosition , 0, QTableWidgetItem(str(produit)))
			self.tableWidget.setItem(self.rowPosition , 1, QTableWidgetItem(str(self.prix)))
			self.tableWidget.setItem(self.rowPosition , 2, QTableWidgetItem(str(self.quantite)))
			self.tableWidget.setItem(self.rowPosition , 3, QTableWidgetItem(str(self.total)))
			self.tableWidget.setItem(self.rowPosition , 5, QTableWidgetItem(str(self.stock)))
			self.tableWidget.setItem(self.rowPosition , 6, QTableWidgetItem(str(date)))
			
			#DATA BINDING
			#DATA TO Products table
			query.execute("UPDATE Products SET product_stock = {0} where product_name = '{1}'".format(int(self.stock)-self.quantite,name))
			conn.commit()
			self.prodQnt.setText("")
            
			# self.tableWidget.cellChanged.connect(self.QChanged) 

            
	# def QChanged(self):#CELL Qnt Changed
        # ==========HOVER ONN TITLE
		# rowPosition = self.tableWidget.rowCount()
		# index = self.tableWidget.currentRow()

		# q = self.tableWidget.item(index,2).text()  
		# p = self.tableWidget.item(index,3).text()  
        
		# sender = self.sender()
		# self.quantite = int(q)
		# print(self.quantite)
			
	def ButtonBC(self):#SLOT : ON CLICK BUTTON BARCODE, SELECTION, SETITEM TO TABLEWIDGET 
		name = self.prodBC.text()
		if name != '':
			try:
				self.produit = self.produit + 1
				self.viewProdNum.setText(str(self.produit))
				
				query.execute('SELECT product_price,product_stock,product_date,product_name,product_BC,product_category FROM Products WHERE product_BC = "{0}"'.format(str(name)))
				fetch=query.fetchall()
				
				prix=fetch[0][0]
				if self.prodQnt.text()!="":
					quantite = int(self.prodQnt.text())
				else:
					quantite= 1
					
				total= prix * quantite
				stock=fetch[0][1]
				date=fetch[0][2]
				produit=fetch[0][3]
				BC=fetch[0][4]
				cat=fetch[0][5]
			
				#APPENDING ON CONTAINERS LISTS
				self.pName.append(produit)
				self.pPrice.append(prix)
				self.pQnt.append(quantite)	
				self.pStock.append(int(stock))
				self.pTotalList.append(int(quantite) * int(prix))
				
				self.a_list.append(total)
				self.viewTotal.display(sum(self.a_list))
				
				#self.tableWidget.setSortingEnabled(False)
				self.rowPosition = self.tableWidget.rowCount()
				self.tableWidget.insertRow(self.rowPosition)
				#self.tableWidget.setSortingEnabled(True)
				self.tableWidget.setItem(self.rowPosition , 0, QTableWidgetItem(str(produit)))
				self.tableWidget.setItem(self.rowPosition , 1, QTableWidgetItem(str(prix)))
				self.tableWidget.setItem(self.rowPosition , 2, QTableWidgetItem(str(quantite)))
				self.tableWidget.setItem(self.rowPosition , 3, QTableWidgetItem(str(total)))
				self.tableWidget.setItem(self.rowPosition , 4, QTableWidgetItem(str(BC)))
				self.tableWidget.setItem(self.rowPosition , 5, QTableWidgetItem(str(stock)))
				self.tableWidget.setItem(self.rowPosition , 6, QTableWidgetItem(str(date)))
				
				query.execute("UPDATE Products SET product_stock = {0} where product_BC = '{1}'".format(int(stock)-int(quantite),name))
				self.prodBC.clear()
				self.prodQnt.setText("")
				conn.commit()
			except:
				self.MF.raiseIndefinedExcept("Produit introuvable !")
				self.viewProdNum.setText((str(int(self.viewProdNum.text()) - 1)))
				self.prodBC.clear()
				if self.prodQnt.text()!="":
					self.prodQnt.setText("")
			
	def OtherProdFunction(self):#SLOT : ON CLICK BUTTON OtherProdManual, SELECTION, SETITEM TO TABLEWIDGET 
		name = self.prodName.text()
		if name != '':
			self.produit = self.produit + 1
			self.viewProdNum.setText(str(self.produit))
			
			produit=name
			prix=self.prodPrice.text()
			if self.prodQnt.text() != "":
				try :
					quantite = int(self.prodQnt.text())
					total= int(prix) * int(quantite)
				except:
					quantite = 1
					total= int(prix) * int(quantite)
			else :
				self.MF.raiseCaseExcept("la case Quantitié")
			stock=""
			date=""
			cat=""
			
			#APPENDING ON CONTAINERS LISTS
			self.pName.append(produit)
			self.pPrice.append(prix)
			self.pQnt.append(quantite)
			self.pStock.append(0)
			self.pTotalList.append(int(quantite) * int(prix))
			
			self.a_list.append(total)
			self.viewTotal.display(sum(self.a_list))
			
			#self.tableWidget.setSortingEnabled(False)
			self.rowPosition = self.tableWidget.rowCount()
			self.tableWidget.insertRow(self.rowPosition)
			#self.tableWidget.setSortingEnabled(True)
			self.tableWidget.setItem(self.rowPosition , 0, QTableWidgetItem(str(produit)))
			self.tableWidget.setItem(self.rowPosition , 1, QTableWidgetItem(str(prix)))
			self.tableWidget.setItem(self.rowPosition , 2, QTableWidgetItem(str(quantite)))
			self.tableWidget.setItem(self.rowPosition , 3, QTableWidgetItem(str(total)))
			self.tableWidget.setItem(self.rowPosition , 5, QTableWidgetItem(str(stock)))
			self.tableWidget.setItem(self.rowPosition , 6, QTableWidgetItem(str(date)))
			
			self.prodName.clear()
			self.prodQnt.setText("")
			self.prodPrice.clear()

#===================================SLOTS
	def ProdEdit(self):#SLOT : ProductCreatorDialog => prodEdit
		dialog=QWidget()
		dialog.ui=self.prodEdit
		dialog.ui.exec_()
		
	def MyProds(self):#SLOT : MyProductDialog => myProd
		self.myProd.MyProdInit()
		dialog=QDialog()
		dialog.ui = self.myProd
		dialog.ui.exec_()
	
#==========RECEIPT 
	def ReiceptFunc(self):#SLOT : Receipt => recette
		self.receipt.MyProdInit()
		dialog=QDialog()
		dialog.ui = self.receipt
		dialog.ui.exec_()
#==========REGISTER 
	def RegisterFunc(self):#SLOT : Receipt => recette
		self.register.registerSum()
		dialog=QDialog()
		dialog.ui = self.register
		dialog.ui.exec_()

	def alerteStockFunc(self):#SLOT : Alert stock
		self.alertSt.InitAlert()
		dialog=QDialog()
		dialog.ui=self.alertSt
		dialog.ui.exec_()

	def alertePeremption(self):#SLOT : Alerts peremption
		self.alertPr.InitAlert()
		dialog=QDialog()
		dialog.ui =self.alertPr
		dialog.ui.exec_()

#==========VALDATION BUTTONS
	def UndoFunction(self):#SLOT : UNDO FUNCTION ON CLICK, ERAESE LAST TABLE WIDGET ROW
		try:
			#==========HOVER ONN TITLE
			rowPosition = self.tableWidget.rowCount()
			value = self.tableWidget.item(rowPosition-1,5).text()
			name = self.tableWidget.item(rowPosition-1,0).text()
			try :
				query.execute("UPDATE Products SET product_stock = {0} where product_name = '{1}'".format(int(value),str(name)))
				conn.commit()
			except:
				pass
			
			self.tableWidget.removeRow(rowPosition-1)
			self.a_list.pop()
			self.viewTotal.display(sum(self.a_list))
			self.produit = self.produit - 1 
			self.viewProdNum.setText(str(self.produit))
			if self.prodQnt.text() != "":
				self.prodQnt.setText("")
			try:
				a=self.received.text()
				c = int(a)-sum(self.a_list)
				if self.discount.text() != "" and self.discount.text() != "0" :
					b=int(self.discount.text())
					c = c + b
				
				self.payBack.setText(str(c))
			except:
				pass
		except:
			pass

	def CancelFunction(self):#SLOT : CANCEL FUNCTION ON CLICK, setRowCount to 0 ON TABLE WIDGET ROW
		
		rowPosition = self.tableWidget.rowCount()
		i=1
		while i <= rowPosition :
			try :
				value = self.tableWidget.item(rowPosition-i,5).text()
				name = self.tableWidget.item(rowPosition-i,0).text()
				query.execute("UPDATE Products SET product_stock = {0} where product_name = '{1}'".format(int(value),str(name)))
			except:
				pass
			query.execute ('DELETE FROM Receipt WHERE product_ref ="{0}" ;'.format( int(self.ref)+1))
			conn.commit()
			i+=1
		
		self.tableWidget.clearContents()
		self.tableWidget.setRowCount(0)
		self.a_list = []
		self.viewTotal.display(0)
		self.produit = 0
		self.viewProdNum.setText(str(self.produit))
		self.TableWidgetInit()
		
		self.received.setText("")
		self.discount.setText("")
		self.payBack.setText("")
		if self.prodQnt.text() != "":
			self.prodQnt.setText("")
			 
	def ValidFunction(self):#SLOT : VALIDATION, DISPLAY MONEY, set TOTAL LIST TO 0, TABLE WIDGET INIT, REF +1,
		if self.viewProdNum.text() != "" and self.viewProdNum.text() != "0":
			self.viewTotal.display(0)
			a=self.received.text()
			if a != '':
				try:
					c = int(a)-sum(self.a_list)
					if self.discount.text() != "" and self.discount.text() != "0" :
						b=int(self.discount.text())
						c = c + b
					
					self.payBack.setText(str(c))
				except:
					pass
			
			self.ref = self.ref + 1
			self.refNum.setText(str(self.ref))
			query.execute('UPDATE ref SET ref ={0} where id=1'.format(self.ref))
			
			if self.prodQnt.text()!="":
				self.prodQnt.setText("")
				
			self.Reference()
			conn.commit()
			
			############ TICKET GENERATION
			try :
				TICKET = Document()
				
				p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")

				h = TICKET.add_heading("\t\t\t\t"+self.society+"", level=1)
				h.bold = True
				h.italic = True
				p = TICKET.add_paragraph("\t\t\t\t"+self.localDateTime)
				
				p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
				tab = TICKET.add_table(1,4)
				heading_cells = tab.rows[0].cells
				heading_cells[0].text = 'Nom produit'
				heading_cells[1].text = 'PRIX Produit'
				heading_cells[2].text = 'Quantité produit'
				heading_cells[3].text = 'Montant Produit'
										
				pi = 0
				for prod in self.pName :
					if self.pQnt[pi] != 0:
						cells = tab.add_row().cells
						cells[0].text = str(self.pName[pi]).strip("(',')")
						cells[1].text = str(self.pPrice[pi]).strip("(',')") + "DA"
						cells[2].text = str(self.pQnt[pi]).strip("(',')")
						cells[3].text = str(self.pTotalList[pi]).strip("(',')") + "DA"
					pi+=1
					
				p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
				try :	
					p = TICKET.add_heading("\t\t\t\tTotal à payer : "+str(sum(self.a_list))+" DA ", level=2)
					p.bold = True
				except :
					return 0
				
				p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
				
				if a != '':
					try:
						if self.discount.text() != "" :
							p = TICKET.add_paragraph("\t\t\t\tReçu : "+self.received.text()+" DA, Rendu : "+str(c)+" DA, Remise : "+str(b)+" DA")
		
						else :	
							p = TICKET.add_paragraph("\t\t\t\tReçu : "+self.received.text()+" DA, Rendu : "+str(c)+" DA, Remise : 0 DA")
					except:
						pass
				p = TICKET.add_paragraph("\t\t\t\tCaissier : Annonyme")
				p = TICKET.add_paragraph("\t\t\t\tNbr de Produits : "+str(self.produit)+".")
				
				p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
				p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
				p = TICKET.add_heading("\t\t\tTicket No : "+str(self.ref)+". --"+self.society+"-- "+self.numero, level=3)
				p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")

				
				TICKET.save('DOCUMENTS/TICKETS/TICKET-'+str(self.ref)+"_"+self.localDate+'.docx' )
				os.startfile('DOCUMENTS\TICKETS\TICKET-'+str(self.ref)+"_"+self.localDate+'.docx' , 'print')
			except :
				pass	
			
			#DATA BINDING ON RECEIPT
			#==========HOVER ONN TITLE
			rowPosition = self.tableWidget.rowCount()
			i=0
			for prod in self.pName :
				query.execute('INSERT INTO Receipt \
			(product_name,product_price,product_qnt,product_date,product_ref)\
			values("{0}",{1} ,{2} ,"{3}",{4});'\
			.format(str(prod), int(self.pPrice[i]) ,int(self.pQnt[i]), str(self.Date),int(self.ref)))
		
				#=============================================REGISTER DEPOT BINDING
				query.execute("INSERT INTO Register_add (add_date, add_depo, add_description,add_value) VALUES ('"\
				+self.Date+"','Vente ','"+str(prod).strip("(',')")+"',"+str(int(int(self.pPrice[i])*int(self.pQnt[i])))+")")
				i+=1
				
				#=============================================DISCOUNTS
			if self.discount.text() != "" and self.discount.text() !='0' :
				query.execute("INSERT INTO Discount (disc_date,disc_ref, disc_value) VALUES ('{0}','{1}','{2}')"\
				.format(self.Date,int(self.ref),int(self.discount.text())))
				#=============================================REGISTER DEPENSE BINDING
				query.execute("INSERT INTO Register_dep (dep_date, dep_type, dep_description,dep_value) VALUES ('"\
				+self.Date+"','Remise','Remise ticket :"+str(self.ref)+"','"+self.discount.text()+"')")	
				
			conn.commit()
				
			self.register.registerSum()
				
			self.a_list=[]
			self.pName = []#container List : selected products name
			self.pPrice = []#container List : selected products price
			self.pQnt = []#container List : selected products quantity
			self.pStock = []#container List : selected products quantity
			self.pTotalList = []#container List : selected products quantity  
			
			self.received.setText("")
			self.discount.setText("")
			self.payBack.setText("")
				
			self.tableWidget.setRowCount(0)
			self.produit=0
			self.viewProdNum.setText(str(self.produit))
	
	def ChangeSlot(self):
		sender = self.sender()
		if sender.text() != "0" :
			try:
				a=self.received.text()
				c = int(a)-sum(self.a_list)
				if self.discount.text() != "" and self.discount.text() != "0" :
					b=int(self.discount.text())
					c = c + b
				
				self.payBack.setText(str(c))
			except:
				pass
			
#============================================================================================================HOMEPAGE 			
qtHome= "DESIGN/WIDGETS/HomeWidget.ui"
Ui_HomeWidget, QtBaseClass = uic.loadUiType(qtHome)
class HomePage(QWidget, Ui_HomeWidget):#CONFIRM : VALID DIALOG ON VLDButton CLICK()
	def __init__(self):
		QWidget.__init__(self)
		Ui_HomeWidget.__init__(self)
		self.setupUi(self)
		
		self.declaredUi()
	def declaredUi(self):
		self.sellPoint = ProductWidget()
		self.connecter = ConnectDialog()
		self.settings = SettingCreatorDialog()
				
qtCreatorFile = "DESIGN/WIDGETS/mainWindow.ui" # Enter file here.
Ui_MainWindow, QtBaseClass = uic.loadUiType(qtCreatorFile)
class MainWindow(QMainWindow, Ui_MainWindow):
	def __init__(self):
		QMainWindow.__init__(self)
		Ui_MainWindow.__init__(self)
		self.setupUi(self)
		
		self.setWindowTitle("My Manager Solution.")
		
		self.MF = MessageFactory()
		
		self.declaredUi()
		self.initUi()
		self.homeSignals()
		
		self.Date = time.strftime("%Y-%m-%d")
		
	def declaredUi(self):
		self.home = HomePage()
		self.sellPoint = ProductWidget()
		self.connecter = ConnectDialog()
		self.register = RegisterCreatorDialog()
		self.settings = SettingCreatorDialog()
		self.dev = DevDialog()
				
	def initUi(self):
		self.dockWidget.setWidget(self.home)
		self.setCentralWidget(self.dockWidget)
		
	def homeSignals(self):
		
		self.home.feature_1.clicked.connect(self.homeSlots)
		self.home.feature_2.clicked.connect(self.homeSlots)
		self.home.feature_3.clicked.connect(self.homeSlots)
		self.home.feature_4.clicked.connect(self.homeSlots)
		
		self.home.feature_5.clicked.connect(self.homeSlots)
		self.home.feature_6.clicked.connect(self.homeSlots)
		self.home.feature_7.clicked.connect(self.homeSlots)
		self.home.feature_8.clicked.connect(self.homeSlots)
		
		#PONTSELL SIGNALS
		self.sellPoint.homePage.clicked.connect(self.homeSlots)
		
		#SETTINGS SIGNALS
		self.settings.setOK.clicked.connect(self.homeSlots)
		self.settings.setAnnuler.clicked.connect(self.homeSlots)
		self.settings.logDel.clicked.connect(self.homeSlots)
		self.settings.disconnect.clicked.connect(self.homeSlots)
		
		#CONNECTER SIGNALS
		self.connecter.connOk.clicked.connect(self.homeSlots)
		
		#DEV SIGNALS
		self.dev.back.clicked.connect(self.homeSlots)
		
	def homeSlots(self) :
		sender = self.sender()
		
		if sender == self.home.feature_1 :
			self.dockWidget.setWidget(self.sellPoint)
		
		if sender == self.home.feature_2 :
			self.register.registerSum()
			dialog=QDialog()
			dialog.ui = self.register
			dialog.ui.exec_()
		
		if sender == self.home.feature_3:
			self.dockWidget.setWidget(self.dev)
		
		if sender == self.home.feature_4:
			self.dockWidget.setWidget(self.settings)
		
		if sender == self.home.feature_5 or  sender == self.home.feature_6 or  sender == self.home.feature_7 or  sender == self.home.feature_8 :
			self.MF.raiseNoPack()
		#POINT SELL SLOTS
		if sender == self.sellPoint.homePage:
			self.dockWidget.setWidget(self.home)
			
		#SETTINGS SLOTS
		if sender == self.settings.setOK or sender == self.settings.setAnnuler :
			self.dockWidget.setWidget(self.home)
			self.settings.selection()
			
		if sender == self.settings.logDel and self.settings.who.text() != "admin" and self.settings.testingco != 0:
			self.connecter.show()
			self.close()
			
		if sender == self.settings.disconnect :
			self.connecter.show()
			self.close()
			
		if sender == self.connecter.connOk :
			self.connecter.login(self)
			self.settings.selection()
			
		if sender == self.dev.back :
			self.dockWidget.setWidget(self.home)
		
	def toolBarCalls(self):	
		if self.sender() == self.HomeAct :
			self.dockWidget.setWidget(self.sellPoint)	
		if self.sender() == self.CaisseAct :
			self.register.totalRegister()
			self.register.show()	
		if self.sender() == self.EditAct :
			self.settings.selection()
			self.connecter.show()
		
class Opener(QDialog):
	def __init__(self):
		QDialog.__init__(self)
		
		self.declaredUi()
		
	def declaredUi(self):
		self.connecter = ConnectDialog()
		self.window = MainWindow()	
		self.home = HomePage()
		
		self.connecter.show()
		self.connecter.connOk.clicked.connect(self.loginSlot)
		
	def loginSlot(self):
		if self.sender() == self.connecter.connOk :
			self.connecter.login(self.window)#
#============================================================================================================END OPENER
	
if __name__ == "__main__":
	app = QApplication(sys.argv)
	opener = Opener()
	window = opener.connecter
	
	MAC=""
	MAC = uuid.UUID(int=uuid.getnode())
	MAC=str(MAC)
	# create_connection(r"db.db")
	
	query.execute("SELECT Key FROM MAC WHERE id=1")
	KEY = query.fetchone()
	tester = str(KEY).strip("(',')")

	if tester == '':
		query.execute("UPDATE MAC SET Key = '%s' WHERE id=1" %MAC)
		conn.commit()
		window.show()
		sys.exit(app.exec_())
	
	elif tester == str(MAC):
		window.show()
		sys.exit(app.exec_())
		
	elif tester != str(MAC):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)

		msg.setText("Vous êtes en présence d'une version du produit obselete")
		msg.setInformativeText("La clé de sécurité est incompatible")
		msg.setWindowTitle("ALERT COPY OBSELETE !")
		msg.exec_()			


